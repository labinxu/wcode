from docx import Document
def write(docname):
    document = Document()
    styles = document.styles
    document.add_heading('Document Tile',0)
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True

    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    #document.add_paragraph('Intense quote', style=style)
    p = document.add_paragraph(
        '', style='ListBullet')
    p.add_hyperlink(text='foobar', url='http://github.com')
    document.add_paragraph(
        'first item in ordered list', style='ListNumber')
    document.add_page_break()
    document.save(docname)
