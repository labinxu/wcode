# -*- coding:utf-8 -*-
import utils
from docx import Document
# #######################
INPUT_ARGS = ''

def debugWritefile(data,filename):
    with open(filename,'w') as f:
        f.write(data)

def parse(filename):
    result = {}
    with open(filename) as f:
        begin = False
        for l in f.readlines():
            r = l.split(INPUT_ARGS.seperator)

            if(not begin):
                if(r[0].strip()=='Module'):
                   begin = True
                continue
            else:
                # status:
                r = r[0:11]
                try:
                    status = r[9].strip()
                    if(status == ''):
                        continue
                    else:
                       if not result.has_key(status):
                           result[status] = {}
                    component = r[8].strip()
                    if component != '':
                        # jiraid, failed number, comments
                        if result[status].has_key(component):
                            result[status][component].append((r[7],r[6],r[-1].strip()))
                        else:
                            result[status][component] = []
                            result[status][component].append((r[7],r[6],r[-1].strip()))

                except IndexError,e:
                    pass

        return result

def output(result):
    document = Document()
    document.add_heading('Document Tile',0)
    for status, components in result.items():
        failures = sumFailures(components)
        if failures <= 0:
            continue

        p = document.add_paragraph( "%s %s.\n" % (status, sumFailures(components)) ,style='ListBullet')
        
        for  component, contents in components.items():

            p = document.add_paragraph( "      %s %s.\n" % (component,sumComponentFailuress(contents)),
                                        style='ListBullet')
            for id, fn, co in contents:
                p.add_hyperlink(text='        %s'%utils.GetJiraId(id),url='https://github.com')
                p.add_run(text="        %s.%s\n" % (fn, co))

    document.add_page_break()
    document.save(INPUT_ARGS.output_file)

def sumComponentFailuress(contents):
    return reduce(lambda x,y:x+y ,[int(i[1]) for i in contents])

def sumFailures(components):
    sumFailures = 0
    for  component, contents in components.items():
        failedNumber = reduce(lambda x,y:x+y,[int(i[1]) for i in contents],0)
        sumFailures += failedNumber

    return sumFailures


def cmdline(args=None):
    import argparse
    global INPUT_ARGS

    parser = argparse.ArgumentParser()
    if(args != None):
        return parser.parse_args(args)

    parser.add_argument('-s', '--seperator', action="store",
                        dest='seperator',
                        default='|',
                        help='seperator for source file')

    parser.add_argument('-o', '--output', action="store",
                        dest='output_file',
                        default='cts_result_st.txt',
                        help='output file')

    parser.add_argument('-i', '--input', action='store',
                        dest='input_file',
                        required=True,
                        default='./data/cts_report.txt',
                        help='input file')

    INPUT_ARGS = parser.parse_args()

    return INPUT_ARGS;


def main():
    args = cmdline()
    result = parse(args.input_file)
    output(result)

if __name__=='__main__':
    main()
